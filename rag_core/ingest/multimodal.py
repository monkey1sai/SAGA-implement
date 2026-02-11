"""多模態文件處理器"""

from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

from .base import DocumentLoader, Document

logger = logging.getLogger(__name__)


class MultimodalIngestor(DocumentLoader):
    """
    多模態文件處理器
    
    統一入口處理各種文件格式：
    - PDF (PyMuPDF)
    - DOCX (python-docx)
    - TXT/MD (純文字)
    
    支援語義分塊，將長文件切分為適合檢索的片段。
    
    Args:
        chunk_size: 分塊大小（字元數）
        chunk_overlap: 分塊重疊（字元數）
    """
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    @property
    def supported_extensions(self) -> List[str]:
        return [".pdf", ".docx", ".txt", ".md", ".markdown"]
    
    async def load(self, file_path: Path) -> List[Document]:
        """載入並處理文件"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix == ".pdf":
            return await self._load_pdf(file_path)
        elif suffix == ".docx":
            return await self._load_docx(file_path)
        elif suffix in [".txt", ".md", ".markdown"]:
            return await self._load_text(file_path)
        else:
            raise ValueError(f"不支援的文件格式: {suffix}")
    
    async def _load_pdf(self, file_path: Path) -> List[Document]:
        """載入 PDF 文件"""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("請安裝 PyMuPDF: pip install pymupdf")
        
        documents = []
        doc = fitz.open(file_path)
        
        full_text = ""
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text()
            full_text += text + "\n"
        
        doc.close()
        
        # 分塊
        chunks = self._chunk_text(full_text)
        
        for i, chunk in enumerate(chunks):
            documents.append(Document(
                text=chunk,
                metadata={
                    "source": str(file_path),
                    "filename": file_path.name,
                    "chunk_index": i,
                    "file_type": "pdf",
                },
                chunk_type="semantic",
            ))
        
        logger.info(f"PDF 載入完成: {file_path.name}, {len(documents)} 個片段")
        return documents
    
    async def _load_docx(self, file_path: Path) -> List[Document]:
        """載入 DOCX 文件"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("請安裝 python-docx: pip install python-docx")
        
        doc = DocxDocument(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        
        # 分塊
        chunks = self._chunk_text(full_text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            documents.append(Document(
                text=chunk,
                metadata={
                    "source": str(file_path),
                    "filename": file_path.name,
                    "chunk_index": i,
                    "file_type": "docx",
                },
                chunk_type="semantic",
            ))
        
        logger.info(f"DOCX 載入完成: {file_path.name}, {len(documents)} 個片段")
        return documents
    
    async def _load_text(self, file_path: Path) -> List[Document]:
        """載入純文字文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            full_text = f.read()
        
        # 分塊
        chunks = self._chunk_text(full_text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            documents.append(Document(
                text=chunk,
                metadata={
                    "source": str(file_path),
                    "filename": file_path.name,
                    "chunk_index": i,
                    "file_type": "text",
                },
                chunk_type="semantic",
            ))
        
        logger.info(f"文字檔載入完成: {file_path.name}, {len(documents)} 個片段")
        return documents
    
    def _chunk_text(self, text: str) -> List[str]:
        """
        將文字分塊
        
        使用滑動窗口方式，確保片段間有重疊以保持語義連貫。
        """
        if not text.strip():
            return []
        
        chunks = []
        start = 0
        text_len = len(text)
        
        while start < text_len:
            end = min(start + self.chunk_size, text_len)
            chunk = text[start:end].strip()
            
            if chunk:
                chunks.append(chunk)
            
            # 下一個起點（考慮重疊）
            start = end - self.chunk_overlap
            if start >= text_len:
                break
        
        return chunks
    
    async def ingest_directory(
        self,
        directory: Path,
        recursive: bool = True,
    ) -> List[Document]:
        """
        批次處理目錄下的所有文件
        
        Args:
            directory: 目錄路徑
            recursive: 是否遞迴處理子目錄
            
        Returns:
            所有文件的片段列表
        """
        directory = Path(directory)
        all_documents = []
        
        pattern = "**/*" if recursive else "*"
        
        for file_path in directory.glob(pattern):
            if file_path.is_file() and self.can_load(file_path):
                try:
                    docs = await self.load(file_path)
                    all_documents.extend(docs)
                except Exception as e:
                    logger.error(f"處理文件失敗 {file_path}: {e}")
        
        logger.info(f"目錄處理完成: {len(all_documents)} 個片段")
        return all_documents
